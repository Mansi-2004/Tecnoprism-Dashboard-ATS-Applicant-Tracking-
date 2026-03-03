"""
Production-grade resume text extraction with OCR fallback.
Extracts and normalizes resume text from PDF/DOCX files.
"""

import io
import re
import pdfplumber
import docx
from fastapi import UploadFile
from typing import Tuple

try:
    from pdf2image import convert_from_bytes
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """
    Extract text from PDF or DOCX file bytes.
    This is the main extraction function that works with raw bytes.
    """
    filename_lower = filename.lower()
    file_ext = filename_lower.split(".")[-1] if "." in filename_lower else ""

    extracted_text = ""

    # PDF handling
    if file_ext == "pdf":
        extracted_text = _extract_from_pdf(content)
        
        # OCR fallback if text is too short
        if OCR_AVAILABLE and (not extracted_text or len(extracted_text.strip()) < 100):
            extracted_text = _extract_from_pdf_ocr(content)
            if not extracted_text:
                raise ValueError("Failed to extract text from PDF (even with OCR)")
        elif not extracted_text:
            raise ValueError("Failed to extract text from PDF")

    # DOCX handling
    elif file_ext == "docx":
        try:
            doc = docx.Document(io.BytesIO(content))
            extracted_text = "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            raise ValueError("Invalid DOCX file. Please convert to .docx if needed.")

    else:
        raise ValueError(f"Unsupported file format: {file_ext}. Please upload a PDF or DOCX file.")

    # Normalize text
    normalized_text = _normalize_text(extracted_text)
    
    if not normalized_text or len(normalized_text.strip()) < 50:
        raise ValueError(f"Resume text is too short or empty after extraction. Got {len(normalized_text)} chars.")
    
    return normalized_text


async def extract_text_from_file(file: UploadFile) -> str:
    """
    Extract text from PDF or DOCX file.
    Falls back to OCR if PDF text extraction is empty.
    Normalizes and returns clean text.
    """
    content = await file.read()
    filename = (file.filename or "").lower()
    file_ext = filename.split(".")[-1] if "." in filename else ""
    content_type = (file.content_type or "").lower()

    extracted_text = ""

    # PDF handling
    if file_ext == "pdf" or "pdf" in content_type:
        extracted_text = _extract_from_pdf(content)
        
        # OCR fallback if text is too short
        if OCR_AVAILABLE and (not extracted_text or len(extracted_text.strip()) < 100):
            extracted_text = _extract_from_pdf_ocr(content)
            if not extracted_text:
                raise ValueError("Failed to extract text from PDF (even with OCR)")
        elif not extracted_text:
            raise ValueError("Failed to extract text from PDF")

    # DOCX handling
    elif file_ext == "docx" or "wordprocessingml" in content_type or "docx" in content_type:
        try:
            doc = docx.Document(io.BytesIO(content))
            extracted_text = "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            raise ValueError("Invalid DOCX file. Please convert to .docx if needed.")

    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")

    # Normalize text
    normalized_text = _normalize_text(extracted_text)
    
    if not normalized_text or len(normalized_text.strip()) < 50:
        raise ValueError("Resume text is too short or empty after extraction.")
    
    return normalized_text


def _extract_from_pdf(content: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        text = ""
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
        return text
    except Exception:
        return ""


def _extract_from_pdf_ocr(content: bytes) -> str:
    """Extract text from PDF using OCR (pytesseract + pdf2image)."""
    if not OCR_AVAILABLE:
        return ""
    
    try:
        images = convert_from_bytes(content)
        text = ""
        for image in images:
            page_text = pytesseract.image_to_string(image)
            text += page_text + "\n"
        return text
    except Exception:
        return ""


def _normalize_text(text: str) -> str:
    """
    Normalize resume text:
    - Collapse horizontal whitespace
    - Preserve newlines for better feature extraction (names, etc.)
    """
    # Collapse horizontal whitespace
    text = re.sub(r'[ \t\f\v]+', ' ', text)
    # Collapse multiple newlines into double newlines
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()


def extract_candidate_info(resume_text: str) -> dict:
    """
    Extract structured candidate information from resume text.
    Returns: {
        name, email, phone, linkedin_url, github_url,
        skills, experience_years, education, projects
    }
    """
    info = {
        "name": None,
        "email": None,
        "phone": None,
        "linkedin_url": None,
        "github_url": None,
        "skills": [],
        "experience_years": 0,
        "education": [],
        "projects": []
    }
    
    text = resume_text.lower()
    
    # Email extraction (case-insensitive)
    email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{1,}\b', resume_text)
    if email_match:
        info["email"] = email_match.group()
    
    # Name extraction - use original case resume_text
    # Most names are at the beginning followed by contact info
    lines = [line.strip() for line in resume_text.split('\n') if line.strip()]
    if not lines and resume_text:
        # If text was flattened into a single line
        lines = [resume_text.strip()[:200]]

    # Better Name extraction logic
    candidate_name = None
    
    # Approach 1: First few lines of the document
    for line in lines[:5]:
        # Skip common resume headers
        if any(keyword in line.lower() for keyword in ['resume', 'curriculum vitae', 'cv', 'profile', 'contact', 'summary']):
            continue
            
        # A name should be 2-4 words, no digits, no common email/phone chars
        words = line.split()
        if 2 <= len(words) <= 5 and not any(char.isdigit() for char in line):
            if not any(char in line for char in ['@', ':', '/', '+', '(', ')']):
                candidate_name = line
                break
                
    # Approach 2: If first approach failed, try the first few words of the entire text
    if not candidate_name:
        words = resume_text.split()
        potential_name = " ".join(words[:3])
        if not any(char.isdigit() for char in potential_name) and '@' not in potential_name:
            candidate_name = potential_name

    if candidate_name:
        # Title case if it's all lower or all upper
        if candidate_name.islower() or candidate_name.isupper():
            info["name"] = candidate_name.title()
        else:
            info["name"] = candidate_name

    # Check if we accidentally got a common keyword as name
    if info["name"] and any(k in info["name"].lower() for k in ['experience', 'education', 'skills', 'objective']):
        info["name"] = None
    
    # Phone extraction (common patterns)
    phone_patterns = [
        r'\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
        r'\(\d{3}\)\s\d{3}-\d{4}',
        r'\d{3}-\d{3}-\d{4}'
    ]
    for pattern in phone_patterns:
        phone_match = re.search(pattern, resume_text)
        if phone_match:
            info["phone"] = phone_match.group()
            break
    
    # LinkedIn URL extraction
    linkedin_match = re.search(r'linkedin\.com/in/[^\s]+', text)
    if linkedin_match:
        info["linkedin_url"] = linkedin_match.group()
    
    # GitHub URL extraction
    github_match = re.search(r'github\.com/[^\s]+', text)
    if github_match:
        info["github_url"] = github_match.group()
    
    # Years of experience extraction
    # First try explicit "X years" patterns
    years_patterns = [
        r'(\d+)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?(?:experience|exp)',
        r'(?:experience|exp)[:\s]+(\d+)\+?\s*(?:years?|yrs?)',
        r'total\s+(?:experience|exp)[:\s]+(\d+)',
    ]
    
    experience_years = 0
    for pattern in years_patterns:
        matches = re.findall(pattern, text)
        if matches:
            years_found = [int(m) for m in matches if m.strip()]
            if years_found:
                experience_years = max(years_found)
                break
    
    # If no explicit years found, try to calculate from date ranges
    # Handles both "May 2024 – Sep 2024" and "may2024–sep2024" (after normalization)
    if experience_years == 0:
        from datetime import datetime
        
        # Pattern to find date ranges: month+year patterns with separators
        # Matches: may2024–sep2024, may2024 - sep2024, May 2024 – Sep 2024, etc.
        date_range_pattern = r'([a-z]{3})\s*(\d{4})\s*[–\-–—]\s*([a-z]{3}|present|current)\s*(\d{4})?'
        date_matches = re.findall(date_range_pattern, text, re.IGNORECASE)
        
        if date_matches:
            months = {
                'jan':1, 'feb':2, 'mar':3, 'apr':4, 'may':5, 'jun':6,
                'jul':7, 'aug':8, 'sep':9, 'oct':10, 'nov':11, 'dec':12
            }
            total_months = 0
            current_dt = datetime.now()
            
            for start_month, start_year, end_month, end_year in date_matches:
                start_month_lower = start_month.lower()[:3]
                end_month_lower = end_month.lower()[:3]
                start_year = int(start_year)
                
                if start_month_lower in months:
                    try:
                        if end_month_lower in months:
                            # Normal date range with explicit end year
                            end_year = int(end_year) if end_year else start_year
                            # If end month < start month, likely next year
                            if months[end_month_lower] < months[start_month_lower]:
                                end_year += 1
                            
                            # Calculate months between dates
                            calc_months = (end_year - start_year) * 12 + (months[end_month_lower] - months[start_month_lower])
                        else:
                            # Open-ended (present/current)
                            calc_months = (current_dt.year - start_year) * 12 + (current_dt.month - months[start_month_lower])
                        
                        if calc_months >= 0:
                            total_months += calc_months
                    except (ValueError, KeyError):
                        continue
            
            # Convert months to years
            if total_months > 0:
                # Round to nearest 0.1 (1 month = 0.08 years approximately)
                experience_years = round(total_months / 12, 1)
                # For very short internships, at least show 0.5 years if there's any duration
                if experience_years < 0.5 and total_months > 0:
                    experience_years = round(total_months / 12, 2)
    
    info["experience_years"] = experience_years
    
    # Education degree extraction (prefer full names over abbreviations)
    # Full degree names mapped to their usual abbreviations
    degree_mapping = {
        'bachelor': ['bachelor', 'baccalaureate', 'b.a.', 'ba', 'b.s.', 'bs', 'b.tech', 'btech'],
        'master': ['master', 'masters', 'm.a.', 'ma', 'm.s.', 'ms', 'm.tech', 'mtech', 'mca', 'mba'],
        'phd': ['phd', 'ph.d.', 'doctorate', 'doctoral'],
        'diploma': ['diploma', 'associate'],
    }
    
    education_found = set()
    for full_name, abbreviations in degree_mapping.items():
        for keyword in [full_name] + abbreviations:
            if keyword in text:
                education_found.add(full_name.title())
                break  # Only count the full name once
    
    info["education"] = list(education_found)
    
    # Skills extraction (look for skill-like words)
    skill_keywords = [
        'python', 'javascript', 'java', 'c++', 'c#', 'rust', 'go', 'kotlin', 'swift',
        'react', 'angular', 'vue', 'node', 'express', 'django', 'flask', 'fastapi',
        'mongodb', 'postgresql', 'mysql', 'redis', 'elasticsearch',
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins',
        'html', 'css', 'sql', 'git', 'linux', 'windows',
        'machine learning', 'nlp', 'computer vision', 'deep learning',
        'rest api', 'graphql', 'websocket', 'microservices', 'devops',
        'agile', 'scrum', 'jira', 'confluence', 'slack'
    ]
    for skill in skill_keywords:
        if skill in text:
            info["skills"].append(skill.title())
    info["skills"] = list(set(info["skills"]))
    
    return info
